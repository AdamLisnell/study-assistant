"""Document parsing for multiple file formats."""

from pathlib import Path
from typing import Optional
import logging

try:
    import pypdf
    PYPDF_AVAILABLE = True
except ImportError:
    PYPDF_AVAILABLE = False

try:
    import pdfplumber
    PDFPLUMBER_AVAILABLE = True
except ImportError:
    PDFPLUMBER_AVAILABLE = False

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

from .utils.logger import setup_logger

logger = setup_logger(__name__)


class DocumentParser:
    """Parse text from various document formats."""
    
    @staticmethod
    def parse_pdf_pypdf(filepath: Path) -> str:
        """Parse PDF using pypdf (faster, simpler)."""
        if not PYPDF_AVAILABLE:
            raise ImportError("pypdf not installed. Run: pip install pypdf")
        
        text = []
        with open(filepath, 'rb') as file:
            pdf_reader = pypdf.PdfReader(file)
            for page_num, page in enumerate(pdf_reader.pages):
                page_text = page.extract_text()
                if page_text:
                    text.append(page_text)
                logger.debug(f"Extracted page {page_num + 1}/{len(pdf_reader.pages)}")
        
        return "\n\n".join(text)
    
    @staticmethod
    def parse_pdf_pdfplumber(filepath: Path) -> str:
        """Parse PDF using pdfplumber (better for complex layouts)."""
        if not PDFPLUMBER_AVAILABLE:
            raise ImportError("pdfplumber not installed. Run: pip install pdfplumber")
        
        text = []
        with pdfplumber.open(filepath) as pdf:
            for page_num, page in enumerate(pdf.pages):
                page_text = page.extract_text()
                if page_text:
                    text.append(page_text)
                logger.debug(f"Extracted page {page_num + 1}/{len(pdf.pages)}")
        
        return "\n\n".join(text)
    
    @staticmethod
    def parse_pdf(filepath: Path, prefer_pdfplumber: bool = True) -> str:
        """
        Parse PDF file to text.
        
        Args:
            filepath: Path to PDF file
            prefer_pdfplumber: Use pdfplumber if available (better quality)
        
        Returns:
            Extracted text
        """
        logger.info(f"Parsing PDF: {filepath.name}")
        
        try:
            if prefer_pdfplumber and PDFPLUMBER_AVAILABLE:
                return DocumentParser.parse_pdf_pdfplumber(filepath)
            elif PYPDF_AVAILABLE:
                return DocumentParser.parse_pdf_pypdf(filepath)
            else:
                raise ImportError("No PDF library available. Install pypdf or pdfplumber")
        except Exception as e:
            logger.error(f"Error parsing PDF {filepath}: {e}")
            raise
    
    @staticmethod
    def parse_docx(filepath: Path) -> str:
        """
        Parse Word document to text.
        
        Args:
            filepath: Path to .docx file
        
        Returns:
            Extracted text
        """
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx not installed. Run: pip install python-docx")
        
        logger.info(f"Parsing Word document: {filepath.name}")
        
        try:
            doc = Document(filepath)
            text = []
            
            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text.append(para.text)
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        text.append(row_text)
            
            logger.debug(f"Extracted {len(text)} paragraphs/rows from Word document")
            return "\n\n".join(text)
            
        except Exception as e:
            logger.error(f"Error parsing Word document {filepath}: {e}")
            raise
    
    @staticmethod
    def parse_text(filepath: Path) -> str:
        """
        Parse plain text file.
        
        Args:
            filepath: Path to text file
        
        Returns:
            File content
        """
        logger.info(f"Reading text file: {filepath.name}")
        return filepath.read_text(encoding="utf-8")
    
    @staticmethod
    def parse_file(filepath: Path) -> str:
        """
        Automatically detect and parse file based on extension.
        
        Args:
            filepath: Path to file
        
        Returns:
            Extracted text content
        
        Raises:
            ValueError: If file format is not supported
        """
        extension = filepath.suffix.lower()
        
        if extension == ".pdf":
            return DocumentParser.parse_pdf(filepath)
        elif extension == ".docx":
            return DocumentParser.parse_docx(filepath)
        elif extension in [".txt", ".md", ".markdown"]:
            return DocumentParser.parse_text(filepath)
        else:
            raise ValueError(f"Unsupported file format: {extension}")
